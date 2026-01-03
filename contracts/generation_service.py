"""
Contract generation service with template population and clause assembly
"""
import io
import hashlib
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .models import (
    Contract, ContractVersion, ContractTemplate, Clause,
    ContractClause, GenerationJob
)


class ContractGenerationService:
    """
    Service for generating contracts from templates and clauses
    """
    
    def __init__(self, tenant_id: str, user_id: str):
        self.tenant_id = tenant_id
        self.user_id = user_id
    
    def generate_contract(
        self,
        template_id: str,
        form_inputs: Dict[str, Any],
        user_instructions: Optional[str] = None,
        async_mode: bool = False
    ) -> Contract:
        """
        Generate a contract from template and inputs
        
        Args:
            template_id: UUID of the template to use
            form_inputs: Structured form data for merge fields
            user_instructions: Optional free-text instructions
            async_mode: Whether to generate asynchronously
            
        Returns:
            Contract instance
        """
        # Load template
        template = ContractTemplate.objects.get(
            id=template_id,
            tenant_id=self.tenant_id,
            status='published'
        )
        
        # Validate business rules
        validation_errors = self._validate_business_rules(template, form_inputs)
        if validation_errors:
            raise ValueError(f"Business rule violations: {validation_errors}")
        
        # Validate mandatory clauses
        missing_clauses = self._check_mandatory_clauses(template)
        if missing_clauses:
            raise ValueError(f"Missing mandatory clauses: {missing_clauses}")
        
        # Create contract record
        contract = Contract.objects.create(
            tenant_id=self.tenant_id,
            template=template,
            title=form_inputs.get('contract_title', 'Untitled Contract'),
            contract_type=template.contract_type,
            created_by=self.user_id,
            counterparty=form_inputs.get('counterparty'),
            value=form_inputs.get('value'),
            start_date=form_inputs.get('start_date'),
            end_date=form_inputs.get('end_date'),
            form_inputs=form_inputs,
            user_instructions=user_instructions,
            status='draft',
            current_version=1
        )
        
        if async_mode:
            # Create async job
            job = GenerationJob.objects.create(
                contract=contract,
                status='pending'
            )
            # TODO: Queue celery task
            return contract
        else:
            # Generate synchronously
            self._generate_document(contract, template, form_inputs)
            return contract
    
    def _generate_document(
        self,
        contract: Contract,
        template: ContractTemplate,
        form_inputs: Dict[str, Any]
    ) -> ContractVersion:
        """
        Generate the actual DOCX document
        """
        # Create new document (in production, load template from R2)
        doc = Document()
        
        # Add title
        title = doc.add_heading(form_inputs.get('contract_title', 'Contract'), 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Contract Type: {template.contract_type}")
        doc.add_paragraph(f"Version: 1")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Populate merge fields
        for field, value in form_inputs.items():
            if value:
                doc.add_paragraph(f"{field.replace('_', ' ').title()}: {value}")
        
        doc.add_paragraph("")
        
        # Assemble clauses
        clauses = self._get_applicable_clauses(template)
        
        doc.add_heading("Contract Clauses", 1)
        
        clause_records = []
        for idx, clause in enumerate(clauses, 1):
            # Add clause
            heading = doc.add_heading(f"{idx}. {clause.name}", 2)
            
            # Add clause metadata (for provenance)
            metadata = doc.add_paragraph(
                f"[Clause ID: {clause.clause_id} | Version: {clause.version}]",
                style='Intense Quote'
            )
            metadata_run = metadata.runs[0]
            metadata_run.font.size = Pt(8)
            metadata_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add clause content
            doc.add_paragraph(clause.content)
            doc.add_paragraph("")
            
            # Track clause usage
            clause_records.append({
                'clause': clause,
                'position': idx,
                'alternatives': self._get_clause_alternatives(clause)
            })
        
        # Save document to memory
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Calculate hash
        doc_content = doc_buffer.getvalue()
        file_hash = hashlib.sha256(doc_content).hexdigest()
        
        # In production: Upload to R2
        r2_key = f"contracts/{contract.tenant_id}/{contract.id}/v1.docx"
        # r2_service.upload(r2_key, doc_buffer)
        
        # Create contract version
        version = ContractVersion.objects.create(
            contract=contract,
            version_number=1,
            r2_key=r2_key,
            template_id=template.id,
            template_version=template.version,
            change_summary="Initial generation",
            created_by=self.user_id,
            file_size=len(doc_content),
            file_hash=file_hash
        )
        
        # Create clause provenance records
        for record in clause_records:
            clause = record['clause']
            ContractClause.objects.create(
                contract_version=version,
                clause_id=clause.clause_id,
                clause_version=clause.version,
                clause_name=clause.name,
                clause_content=clause.content,
                is_mandatory=clause.is_mandatory,
                position=record['position'],
                alternatives_suggested=record['alternatives']
            )
        
        return version
    
    def _validate_business_rules(
        self,
        template: ContractTemplate,
        form_inputs: Dict[str, Any]
    ) -> List[str]:
        """
        Validate business rules from template
        """
        errors = []
        rules = template.business_rules
        
        # Check required fields
        required_fields = rules.get('required_fields', [])
        for field in required_fields:
            if field not in form_inputs or not form_inputs[field]:
                errors.append(f"Missing required field: {field}")
        
        # Check value constraints
        if 'min_value' in rules and 'value' in form_inputs:
            if float(form_inputs['value']) < float(rules['min_value']):
                errors.append(f"Value must be at least {rules['min_value']}")
        
        if 'max_value' in rules and 'value' in form_inputs:
            if float(form_inputs['value']) > float(rules['max_value']):
                errors.append(f"Value must not exceed {rules['max_value']}")
        
        # Check jurisdiction (if required)
        if 'allowed_jurisdictions' in rules and 'jurisdiction' in form_inputs:
            if form_inputs['jurisdiction'] not in rules['allowed_jurisdictions']:
                errors.append(
                    f"Jurisdiction must be one of: {', '.join(rules['allowed_jurisdictions'])}"
                )
        
        return errors
    
    def _check_mandatory_clauses(self, template: ContractTemplate) -> List[str]:
        """
        Check if all mandatory clauses exist
        """
        mandatory_clause_ids = template.mandatory_clauses
        missing = []
        
        for clause_id in mandatory_clause_ids:
            exists = Clause.objects.filter(
                tenant_id=self.tenant_id,
                clause_id=clause_id,
                status='published'
            ).exists()
            
            if not exists:
                missing.append(clause_id)
        
        return missing
    
    def _get_applicable_clauses(self, template: ContractTemplate) -> List[Clause]:
        """
        Get all applicable clauses for this template
        """
        # Get mandatory clauses first
        mandatory_clauses = list(Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=template.mandatory_clauses,
            contract_type=template.contract_type,
            status='published'
        ).order_by('clause_id'))
        
        # Get optional clauses
        optional_clauses = list(Clause.objects.filter(
            tenant_id=self.tenant_id,
            contract_type=template.contract_type,
            is_mandatory=False,
            status='published'
        ).exclude(
            clause_id__in=template.mandatory_clauses
        ).order_by('clause_id'))
        
        # Combine (mandatory first)
        return mandatory_clauses + optional_clauses
    
    def _get_clause_alternatives(self, clause: Clause) -> List[Dict[str, Any]]:
        """
        Get alternative clauses with confidence scores
        """
        alternatives = []
        
        for alt_data in clause.alternatives:
            alt_clause_id = alt_data.get('clause_id')
            rationale = alt_data.get('rationale', '')
            confidence = alt_data.get('confidence', 0.5)
            
            # Try to get the alternative clause
            try:
                alt_clause = Clause.objects.get(
                    tenant_id=self.tenant_id,
                    clause_id=alt_clause_id,
                    status='published'
                )
                
                alternatives.append({
                    'clause_id': alt_clause_id,
                    'clause_name': alt_clause.name,
                    'rationale': rationale,
                    'confidence': confidence
                })
            except Clause.DoesNotExist:
                continue
        
        return alternatives
    
    def approve_contract(self, contract_id: str) -> Contract:
        """
        Approve a contract for download/send/sign
        """
        contract = Contract.objects.get(
            id=contract_id,
            tenant_id=self.tenant_id
        )
        
        contract.is_approved = True
        contract.approved_by = self.user_id
        contract.approved_at = datetime.now()
        contract.status = 'approved'
        contract.save()
        
        return contract
    
    def create_new_version(
        self,
        contract_id: str,
        changes: Dict[str, Any],
        change_summary: str
    ) -> ContractVersion:
        """
        Create a new contract version (never overwrite)
        """
        contract = Contract.objects.get(
            id=contract_id,
            tenant_id=self.tenant_id
        )
        
        # Increment version
        new_version_number = contract.current_version + 1
        
        # Merge changes with existing inputs
        updated_inputs = {**contract.form_inputs, **changes}
        
        # Update contract
        contract.form_inputs = updated_inputs
        contract.current_version = new_version_number
        contract.is_approved = False  # Reset approval
        contract.save()
        
        # Regenerate document
        template = contract.template
        version = self._generate_document(contract, template, updated_inputs)
        version.version_number = new_version_number
        version.change_summary = change_summary
        version.save()
        
        return version
