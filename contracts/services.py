"""
Contract Generation Service

Handles template-based contract generation with:
- DOCX template processing
- Merge field replacement
- Clause assembly
- Version management
- Business rule validation
"""
import uuid
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from django.utils import timezone

from .models import (
    Contract, ContractVersion, ContractClause, 
    ContractTemplate, Clause, BusinessRule, WorkflowLog
)


class RuleEngine:
    """
    Business rule evaluation engine
    """
    
    @staticmethod
    def evaluate_condition(condition: Dict, context: Dict) -> bool:
        """
        Evaluate a single condition against context data
        
        Conditions format:
        {
            "contract_value__gte": 10000000,
            "contract_type": "MSA",
            "start_date__year": 2026
        }
        """
        for key, expected_value in condition.items():
            if '__' in key:
                field, operator = key.rsplit('__', 1)
                actual_value = context.get(field)
                
                if actual_value is None:
                    return False
                
                if operator == 'gte' and not (actual_value >= expected_value):
                    return False
                elif operator == 'lte' and not (actual_value <= expected_value):
                    return False
                elif operator == 'gt' and not (actual_value > expected_value):
                    return False
                elif operator == 'lt' and not (actual_value < expected_value):
                    return False
                elif operator == 'in' and actual_value not in expected_value:
                    return False
                elif operator == 'contains' and expected_value not in actual_value:
                    return False
            else:
                # Direct equality check
                if context.get(key) != expected_value:
                    return False
        
        return True
    
    @classmethod
    def get_mandatory_clauses(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict) -> List[Dict]:
        """
        Get list of mandatory clauses based on business rules
        """
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='mandatory_clause',
            is_active=True
        ).filter(
            models.Q(contract_types=[]) | models.Q(contract_types__contains=[contract_type])
        ).order_by('-priority')
        
        mandatory_clauses = []
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                mandatory_clauses.append({
                    'clause_id': rule.action.get('clause_id'),
                    'message': rule.action.get('message', ''),
                    'rule_name': rule.name
                })
        
        return mandatory_clauses
    
    @classmethod
    def get_clause_suggestions(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict, clause_id: str) -> List[Dict]:
        """
        Get alternative clause suggestions based on rules
        """
        # First get the clause and its predefined alternatives
        try:
            clause = Clause.objects.get(
                tenant_id=tenant_id,
                clause_id=clause_id,
                status='published'
            )
        except Clause.DoesNotExist:
            return []
        
        suggestions = []
        
        # Evaluate predefined alternatives
        for alt in clause.alternatives:
            trigger_rules = alt.get('trigger_rules', {})
            if not trigger_rules or cls.evaluate_condition(trigger_rules, context):
                suggestions.append({
                    'clause_id': alt['clause_id'],
                    'rationale': alt.get('rationale', 'Alternative clause'),
                    'confidence': alt.get('confidence', 0.8),
                    'source': 'predefined'
                })
        
        # Check dynamic business rules for suggestions
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='clause_suggestion',
            is_active=True,
            action__target_clause_id=clause_id
        ).order_by('-priority')
        
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                suggestions.append({
                    'clause_id': rule.action.get('suggest_clause_id'),
                    'rationale': rule.action.get('rationale', rule.description),
                    'confidence': rule.action.get('confidence', 0.7),
                    'source': 'rule_based'
                })
        
        return suggestions
    
    @classmethod
    def validate_contract(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict, selected_clauses: List[str]) -> Tuple[bool, List[str]]:
        """
        Validate contract against business rules
        Returns: (is_valid, error_messages)
        """
        errors = []
        
        # Check mandatory clauses
        mandatory = cls.get_mandatory_clauses(tenant_id, contract_type, context)
        for req in mandatory:
            if req['clause_id'] not in selected_clauses:
                errors.append(f"Mandatory clause missing: {req['clause_id']} - {req['message']}")
        
        # Check validation rules
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='validation',
            is_active=True
        ).filter(
            models.Q(contract_types=[]) | models.Q(contract_types__contains=[contract_type])
        )
        
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                if rule.action.get('type') == 'error':
                    errors.append(rule.action.get('message', rule.description))
        
        return len(errors) == 0, errors


class ContractGenerator:
    """
    Contract generation from templates with clause assembly
    """
    
    def __init__(self, user_id: uuid.UUID, tenant_id: uuid.UUID):
        self.user_id = user_id
        self.tenant_id = tenant_id
        self.rule_engine = RuleEngine()
    
    def generate_from_template(
        self, 
        template_id: uuid.UUID,
        structured_inputs: Dict,
        user_instructions: Optional[str] = None,
        title: Optional[str] = None
    ) -> Contract:
        """
        Generate a new contract from a template
        
        Args:
            template_id: Template to use
            structured_inputs: Form data for merge fields
            user_instructions: Optional free-text instructions
            title: Contract title
        
        Returns:
            Created Contract instance
        """
        template = ContractTemplate.objects.get(
            id=template_id,
            tenant_id=self.tenant_id,
            status='published'
        )
        
        # Create contract record
        contract = Contract.objects.create(
            tenant_id=self.tenant_id,
            template=template,
            title=title or f"{template.contract_type} - {structured_inputs.get('counterparty', 'Draft')}",
            contract_type=template.contract_type,
            counterparty=structured_inputs.get('counterparty'),
            value=structured_inputs.get('value'),
            start_date=structured_inputs.get('start_date'),
            end_date=structured_inputs.get('end_date'),
            form_inputs=structured_inputs,
            user_instructions=user_instructions,
            created_by=self.user_id,
            status='draft',
            current_version=1
        )
        
        # Log creation
        WorkflowLog.objects.create(
            contract=contract,
            action='created',
            performed_by=self.user_id,
            comment=f'Generated from template: {template.name} v{template.version}'
        )
        
        return contract
    
    def create_version(
        self, 
        contract: Contract, 
        selected_clauses: Optional[List[str]] = None,
        change_summary: Optional[str] = None
    ) -> ContractVersion:
        """
        Create a new version of a contract with selected clauses
        
        Args:
            contract: Contract instance
            selected_clauses: List of clause IDs to include (None = use template defaults)
            change_summary: Summary of changes
        
        Returns:
            Created ContractVersion with assembled document
        """
        template = contract.template
        context = {
            'contract_type': contract.contract_type,
            'contract_value': float(contract.value) if contract.value else 0,
            'counterparty': contract.counterparty,
            **contract.form_inputs
        }
        
        # Determine which clauses to include
        if selected_clauses is None:
            # Use template defaults + mandatory clauses
            selected_clauses = list(template.mandatory_clauses)
            mandatory = self.rule_engine.get_mandatory_clauses(
                self.tenant_id, 
                contract.contract_type, 
                context
            )
            for req in mandatory:
                if req['clause_id'] not in selected_clauses:
                    selected_clauses.append(req['clause_id'])
        
        # Validate
        is_valid, errors = self.rule_engine.validate_contract(
            self.tenant_id,
            contract.contract_type,
            context,
            selected_clauses
        )
        
        if not is_valid:
            raise ValidationError({"clauses": errors})
        
        # Create document
        doc = self._create_document(contract, selected_clauses, context)
        
        # Save to BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Calculate hash
        content = doc_bytes.getvalue()
        file_hash = hashlib.sha256(content).hexdigest()
        
        # Create version
        version_number = contract.current_version
        version = ContractVersion.objects.create(
            contract=contract,
            version_number=version_number,
            template_id=template.id,
            template_version=template.version,
            change_summary=change_summary or f'Version {version_number}',
            created_by=self.user_id,
            file_size=len(content),
            file_hash=file_hash,
            r2_key=f'contracts/{contract.id}/v{version_number}.docx'
        )
        
        # Store clause provenance
        self._store_clause_provenance(version, selected_clauses, context)
        
        # Update contract version
        contract.current_version = version_number + 1
        contract.save()
        
        # Log version creation
        WorkflowLog.objects.create(
            contract=contract,
            action='version_created',
            performed_by=self.user_id,
            comment=f'Version {version_number} created',
            metadata={'clause_count': len(selected_clauses)}
        )
        
        # TODO: Upload to R2 storage
        # For now, we'll store the path only
        
        return version
    
    def _create_document(self, contract: Contract, clause_ids: List[str], context: Dict) -> Document:
        """
        Create DOCX document by assembling clauses
        """
        doc = Document()
        
        # Add header
        doc.add_heading(contract.title, 0)
        
        # Add contract metadata
        p = doc.add_paragraph()
        p.add_run(f"Contract Type: ").bold = True
        p.add_run(f"{contract.contract_type}\n")
        p.add_run(f"Counterparty: ").bold = True
        p.add_run(f"{contract.counterparty or 'N/A'}\n")
        if contract.value:
            p.add_run(f"Value: ").bold = True
            p.add_run(f"${contract.value:,.2f}\n")
        p.add_run(f"Created: ").bold = True
        p.add_run(f"{contract.created_at.strftime('%Y-%m-%d')}\n")
        
        doc.add_paragraph()  # Spacing
        
        # Add clauses in order
        clauses = Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=clause_ids,
            status='published'
        ).order_by('clause_id')
        
        for i, clause in enumerate(clauses, 1):
            # Clause heading
            heading = doc.add_heading(f"{i}. {clause.name}", level=2)
            
            # Clause content with merge field replacement
            content = self._replace_merge_fields(clause.content, context)
            doc.add_paragraph(content)
            
            # Add provenance comment (metadata)
            provenance = doc.add_paragraph()
            provenance.add_run(f"[Clause ID: {clause.clause_id} v{clause.version}]").font.size = Pt(8)
            provenance.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
        
        return doc
    
    def _replace_merge_fields(self, text: str, context: Dict) -> str:
        """
        Replace {{field_name}} with actual values
        """
        result = text
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in result:
                result = result.replace(placeholder, str(value))
        return result
    
    def _store_clause_provenance(self, version: ContractVersion, clause_ids: List[str], context: Dict):
        """
        Store clause provenance in contract_clauses table
        """
        clauses = Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=clause_ids,
            status='published'
        ).order_by('clause_id')
        
        for position, clause in enumerate(clauses, 1):
            # Get alternatives for this clause
            alternatives = self.rule_engine.get_clause_suggestions(
                self.tenant_id,
                version.contract.contract_type,
                context,
                clause.clause_id
            )
            
            ContractClause.objects.create(
                contract_version=version,
                clause_id=clause.clause_id,
                clause_version=clause.version,
                clause_name=clause.name,
                clause_content=clause.content,
                is_mandatory=clause.is_mandatory,
                position=position,
                alternatives_suggested=alternatives
            )


from rest_framework.exceptions import ValidationError
from django.db import models
